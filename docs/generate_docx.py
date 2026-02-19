"""PRD markdown → Word 변환 스크립트"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '맑은 고딕'
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 섹션 여백
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_table(headers, rows, col_widths=None):
    """표 추가 헬퍼"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)

# ═══════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI 보행분석 시스템')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('제품 소개 및 기술 문서')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tagline.add_run('스마트폰 영상 한 편으로 14가지 보행 지표를 자동 분석')
run3.font.size = Pt(13)
run3.italic = True

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = info.add_run('기반 검사: 10-Meter Walk Test (10MWT)\n전 세계적으로 가장 널리 사용되는 표준 보행 평가 도구')
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 1. 제품 개요
# ═══════════════════════════════════════════════════
doc.add_heading('1. 제품 개요', level=1)

doc.add_heading('핵심 가치: 기존 방식 vs AI 보행분석', level=2)
add_table(
    ['기존 방식', 'AI 보행분석 시스템'],
    [
        ['스톱워치로 시간만 측정', '14가지 정밀 보행 지표 자동 측정'],
        ['검사자 숙련도에 따라 결과 편차', 'AI가 일관된 기준으로 분석'],
        ['보행속도 1가지만 기록', '속도, 보폭, 활보장, 대칭성, 비율 등'],
        ['고가 장비(압력판, 센서) 필요', '스마트폰 카메라 1대면 충분'],
        ['결과 해석에 전문 지식 필요', '자동 판정 + 임상 코멘트 제공'],
        ['기록 보관 어려움', '환자별 이력 관리 + 추이 그래프'],
    ],
    [3.0, 3.0]
)

# ═══════════════════════════════════════════════════
# 2. 타겟 사용자
# ═══════════════════════════════════════════════════
doc.add_heading('2. 타겟 사용자', level=1)

doc.add_heading('1차: 재활 전문가 / 물리치료사', level=2)
add_bullet('뇌졸중, 파킨슨, 정형외과 수술 후 재활 환자의 보행 평가')
add_bullet('치료 전후 비교를 통한 개선도 객관적 측정')
add_bullet('보험 심사 및 진료 기록에 활용 가능한 정량적 데이터')

doc.add_heading('2차: 요양시설 / 노인복지관', level=2)
add_bullet('고령자 낙상 위험도 스크리닝')
add_bullet('정기적 보행 능력 모니터링')
add_bullet('지역사회 보행 가능 여부(Community Ambulation) 판정')

doc.add_heading('3차: 연구자 / 스포츠 트레이너', level=2)
add_bullet('보행 패턴 연구 데이터 수집')
add_bullet('선수 부상 복귀 시 보행 대칭성 평가')
add_bullet('운동 프로그램 전후 효과 측정')

# ═══════════════════════════════════════════════════
# 3. 핵심 기능 상세
# ═══════════════════════════════════════════════════
doc.add_heading('3. 핵심 기능 상세', level=1)

doc.add_heading('3-1. AI 영상 분석 파이프라인', level=2)
p = doc.add_paragraph()
run = p.add_run('[영상 업로드]  →  [2-Pass 분석]  →  [자동 판정]  →  [리포트 생성]')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_heading('Pass 1: 캘리브레이션 (거리 보정)', level=3)
add_bullet('환자 키(cm) 입력 → AI가 영상 속 키 픽셀 자동 측정', bold_prefix='입력: ')
add_bullet('PPM(Pixels Per Meter) 산출 → 모든 거리 측정의 기준', bold_prefix='결과: ')
add_bullet('카메라와의 거리에 따라 프레임마다 동적 PPM 조정', bold_prefix='원근 보정: ')
add_bullet('고가의 거리 측정 장비 없이도 cm 단위 정밀 거리 측정 가능')

doc.add_heading('Pass 2: 보행 분석', level=3)
add_bullet('YOLOv8-Pose: 실시간 사람 감지 + 17개 관절 (NVIDIA GPU 가속)', bold_prefix='자세 추정 ① ')
add_bullet('MediaPipe Pose Heavy: 33개 관절 정밀 추정 (발꿈치/발끝 포함)', bold_prefix='자세 추정 ② ')
add_bullet('전신이 화면에 보이는 순간 자동 시작/종료', bold_prefix='보행 구간 감지: ')
add_bullet('3가지 독립 알고리즘 → 품질 기반 자동 선택', bold_prefix='Heel Strike 감지: ')
add_bullet('  - Y-peak (발목 수직 위치 피크)')
add_bullet('  - Vx-crossing (발목 속도 전환점)')
add_bullet('  - Foot Separation (양발 간격 변화)')

doc.add_heading('3-2. 14가지 정밀 보행 지표', level=2)

doc.add_heading('보행 지표 (2개)', level=3)
add_table(
    ['지표', '설명', '정상 범위', '임상 의미'],
    [
        ['보행속도', '단위 시간당 이동 거리', '1.0~1.4 m/s', '전반적 보행 능력의 핵심 지표\n("제6의 활력징후")'],
        ['분속수 (Cadence)', '1분당 걸음 수', '100~120 steps/min', '보행 리듬 및 자동화 수준'],
    ],
    [1.3, 1.8, 1.2, 2.0]
)

doc.add_heading('거리 변수 (4개) — 좌/우 개별 측정', level=3)
add_table(
    ['지표', '설명', '정상 범위', '임상 의미'],
    [
        ['좌측 보폭 (L Step)', '오른발→왼발 착지 거리', '55~70 cm', '좌측 하지 추진력'],
        ['우측 보폭 (R Step)', '왼발→오른발 착지 거리', '55~70 cm', '우측 하지 추진력'],
        ['좌측 활보장 (L Stride)', '왼발→왼발 1보행주기 거리', '110~140 cm', '좌측 보행주기 완성도'],
        ['우측 활보장 (R Stride)', '오른발→오른발 1보행주기 거리', '110~140 cm', '우측 보행주기 완성도'],
    ],
    [1.5, 1.8, 1.0, 2.0]
)

doc.add_heading('시간 변수 (4개)', level=3)
add_table(
    ['지표', '설명', '정상 범위', '임상 의미'],
    [
        ['좌측 스텝 시간', '반대발→좌발 착지 소요 시간', '0.50~0.60 s', '좌측 타이밍 조절'],
        ['우측 스텝 시간', '반대발→우발 착지 소요 시간', '0.50~0.60 s', '우측 타이밍 조절'],
        ['좌측 활보 시간', '좌발 1보행주기 소요 시간', '1.00~1.20 s', '좌측 보행 주기'],
        ['우측 활보 시간', '우발 1보행주기 소요 시간', '1.00~1.20 s', '우측 보행 주기'],
    ],
    [1.3, 2.0, 1.0, 2.0]
)

doc.add_heading('비율 변수 (4개)', level=3)
add_table(
    ['지표', '설명', '정상 범위', '임상 의미'],
    [
        ['입각기 비율 (Stance)', '발이 바닥에 접촉한 비율', '58~62%', '안정성 지표'],
        ['유각기 비율 (Swing)', '발이 공중에 있는 비율', '38~42%', '추진력 지표'],
        ['스윙/스탠스 비율', 'Swing ÷ Stance', '~0.67', '보행 효율성'],
        ['좌우 대칭성 (SI)', '좌우 차이 비율', '< 10%', '비대칭 = 병리적 보행'],
    ],
    [1.5, 1.8, 1.0, 2.0]
)

# ── 3-3. 자동 판정 시스템 ──
doc.add_heading('3-3. 자동 판정 시스템', level=2)

doc.add_heading('Perry 보행 속도 분류 (Perry et al., 1995)', level=3)
p = doc.add_paragraph('전 세계적으로 가장 널리 인용되는 보행 속도 분류 체계:')
p.runs[0].font.size = Pt(10)
add_table(
    ['분류', '속도 기준', '임상 의미'],
    [
        ['Household Ambulator', '< 0.4 m/s', '실내 보행만 가능, 높은 의존도'],
        ['Limited Community', '0.4~0.8 m/s', '제한적 외출, 보조 필요'],
        ['Full Community', '0.8~1.0 m/s', '독립적 외출 가능'],
        ['Normal Speed', '1.0~1.4 m/s', '정상 보행 능력'],
    ],
    [1.8, 1.5, 3.0]
)

doc.add_heading('보행 패턴 자동 감지 (3가지)', level=3)
add_table(
    ['패턴', '감지 조건', '임상 의미'],
    [
        ['안정성 보상 패턴', '↑입각기 + ↓유각기 + ↓속도', '균형 불안 시 보이는 보상 전략'],
        ['단보 패턴', '↓보폭 + ↓활보장 + ↑분속수', '한국 고령자에게 흔한 패턴\n(Kim et al., 2024)'],
        ['높은 변동성', 'SI ≥ 10%', '좌우 비대칭 → 편측 약화 의심'],
    ],
    [1.5, 2.0, 2.8]
)

doc.add_heading('대칭성 지수 (Symmetry Index)', level=3)
add_bullet('공식: SI = |좌-우| / (0.5 × (좌+우)) × 100%')
add_bullet('기준: < 10% = 정상 (Patterson et al., 2010; Herzog et al., 1989)')
add_bullet('보폭, 활보장, 스텝시간, 활보시간, 스윙/스탠스 6가지 항목 개별 SI 산출')

# ── 3-4. 근거 영상 클립 ──
doc.add_heading('3-4. 근거 영상 클립 (Evidence Clips)', level=2)
add_bullet('각 측정 변수마다 해당 수치가 측정된 영상 구간을 자동 추출')
add_bullet('평균값에 가장 가까운 대표 스텝을 자동 선택하여 재생')
add_bullet('영상 위에 실시간 오버레이: 발목 추적, HS 마커, 거리/시간 카운터, 판정 게이지')
p = doc.add_paragraph()
run = p.add_run('→ "이 수치가 어떻게 나왔는지" 영상으로 직접 확인 = 기록의 신뢰도 극대화')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ── 3-5. 환자 관리 ──
doc.add_heading('3-5. 환자 관리 및 이력 추적', level=2)
add_bullet('환자 프로필 등록 (이름, 생년월일, 진단명, 키)')
add_bullet('검사 이력 자동 저장 및 조회')
add_bullet('추이 그래프: 시간에 따른 보행속도 변화 시각화', bold_prefix='')
add_bullet('검사 비교: 2개 검사 결과를 나란히 비교 → 치료 효과 객관적 확인')
add_bullet('대시보드: 총 환자 수, 오늘 검사 수, 주간 통계, 진단별 분포')

# ═══════════════════════════════════════════════════
# 4. 기존 방식과의 비교
# ═══════════════════════════════════════════════════
doc.add_heading('4. 기존 방식과의 비교', level=1)

doc.add_heading('4-1. vs 스톱워치 10MWT', level=2)
add_table(
    ['항목', '스톱워치', 'AI 보행분석'],
    [
        ['측정 지표', '보행속도 1가지', '14가지 지표 + 대칭성 + 패턴'],
        ['측정 정밀도', '검사자 반응시간 오차\n(±0.2~0.5s)', '프레임 단위 정밀도\n(±0.017s @60fps)'],
        ['좌우 분리', '불가', '좌/우 개별 측정 + SI'],
        ['재현성', '검사자 간 편차 존재', '동일 영상 → 동일 결과'],
        ['기록 보존', '종이 기록, 분실 위험', '디지털 이력 + 영상 보존'],
        ['근거 확인', '불가 (기억에 의존)', '영상 클립으로 즉시 확인'],
        ['소요 시간', '~5분 (세팅+3회 시행)', '~2분 (촬영1회+자동분석)'],
    ],
    [1.2, 2.3, 2.8]
)

doc.add_heading('4-2. vs 고가 보행분석 장비', level=2)
add_table(
    ['항목', 'GAITRite / Vicon 등', 'AI 보행분석'],
    [
        ['장비 비용', '3,000만원~1억원+', '스마트폰 + PC\n(기존 장비 활용)'],
        ['설치 공간', '전용 보행로 필요\n(6~10m)', '일반 복도에서 촬영 가능'],
        ['측정 지표', '30+ 정밀 지표', '14가지 핵심 임상 지표'],
        ['측정 정밀도', '매우 높음 (mm 단위)', '높음 (cm 단위, 키 기반 보정)'],
        ['접근성', '장비 보유 기관만', '인터넷 연결된 곳이면 어디서나'],
    ],
    [1.2, 2.3, 2.8]
)

doc.add_heading('4-3. 핵심 차별점', level=2)
add_bullet('스마트폰 카메라 1대 + 환자 키 정보만 있으면 시작', bold_prefix='진입 장벽 제로: ')
add_bullet('듀얼 AI 모델 + 동적 원근 보정 + 다중 HS 감지', bold_prefix='임상 수준 정밀도: ')
add_bullet('모든 수치에 대한 영상 증거 클립 자동 생성', bold_prefix='투명한 근거: ')
add_bullet('Perry 분류 + 정상범위 판정 + 패턴 감지 + 임상 코멘트', bold_prefix='즉각적 임상 해석: ')
add_bullet('환자별 시계열 데이터 → 치료 효과의 객관적 입증', bold_prefix='이력 추적: ')

# ═══════════════════════════════════════════════════
# 5. 기술 신뢰성
# ═══════════════════════════════════════════════════
doc.add_heading('5. 기술 신뢰성', level=1)

doc.add_heading('5-1. 사용된 AI 모델', level=2)
add_table(
    ['모델', '용도', '특징'],
    [
        ['YOLOv8n-Pose', '실시간 사람 감지\n+ 17개 관절 추정', 'Ultralytics\nCOCO 학습, GPU 가속'],
        ['MediaPipe Pose Heavy', '33개 관절 정밀 추정', 'Google\n발꿈치/발끝 감지 포함'],
    ],
    [1.8, 2.0, 2.5]
)

doc.add_heading('5-2. 판정 기준 학술 근거', level=2)
add_table(
    ['기준', '출처'],
    [
        ['보행속도 분류', 'Perry J et al. (1995) Classification of Walking Handicap, Stroke'],
        ['대칭성 지수 10% 기준', 'Patterson KK et al. (2010); Herzog W et al. (1989)'],
        ['정상 보행 범위', 'AAFP (2010) Gait Disorders in Older Adults'],
        ['낙상 위험 속도 기준', 'Verghese J et al. (2009) Quantitative gait dysfunction'],
        ['단보 패턴 (한국)', 'Kim YH et al. (2024) Gait characteristics in Korean elderly'],
        ['장애 위험 속도', 'Cesari M et al. (2005) Prognostic value of usual gait speed'],
    ],
    [2.0, 4.3]
)

# ═══════════════════════════════════════════════════
# 6. 사용 워크플로우
# ═══════════════════════════════════════════════════
doc.add_heading('6. 사용 워크플로우', level=1)

steps_data = [
    ('1단계: 검사 설정', '환자 선택 + 키 입력 + 거리 설정 (기본 10m)'),
    ('2단계: 영상 촬영 & 업로드', '스마트폰으로 측면에서 보행 영상 촬영 → 업로드\n촬영 가이드: 측면 90°, 전신 포함, 10m 직선 보행'),
    ('3단계: AI 자동 분석', 'Pass 1: 키 캘리브레이션 (거리 보정)\nPass 2: 보행 분석 (HS 감지, 거리/시간 측정)\n실시간 진행률 표시 + 라이브 미리보기\n소요 시간: 약 1~2분'),
    ('4단계: 결과 확인', 'Perry 보행속도 분류\n14가지 변수 판정표 (정상/초과/미달 + 편차%)\n좌우 대칭성 분석 (6가지 SI)\n보행 패턴 감지 + 임상 코멘트\n영상 클립 근거 확인'),
]
add_table(
    ['단계', '내용'],
    steps_data,
    [1.8, 4.5]
)

# ═══════════════════════════════════════════════════
# 7. 장단점
# ═══════════════════════════════════════════════════
doc.add_heading('7. 장점과 한계 (투명한 공개)', level=1)

doc.add_heading('장점', level=2)
advantages = [
    ('경제성: ', '스마트폰 + PC만으로 임상급 보행분석 가능'),
    ('객관성: ', 'AI 기반 일관된 분석 → 검사자 간 변동성 제거'),
    ('포괄성: ', '1회 촬영으로 14가지 지표 동시 측정'),
    ('투명성: ', '영상 근거 클립으로 모든 수치의 출처 확인 가능'),
    ('편의성: ', '촬영 → 업로드 → 2분 내 결과 완성'),
    ('추적성: ', '환자별 이력 관리 + 추이 그래프로 치료 효과 시각화'),
    ('접근성: ', '1차 의료기관, 요양시설에서도 활용 가능'),
    ('학술 근거: ', '모든 판정 기준에 국제 학술지 참고문헌 명시'),
]
for bold, text in advantages:
    add_bullet(text, bold_prefix=bold)

doc.add_heading('한계 및 보완 계획', level=2)
limits = [
    ('측면 촬영 전용: ', '현재 측면 카메라에서만 정확한 측정 → 향후 다중 앵글 지원 예정'),
    ('거리 정밀도: ', '키 기반 보정은 고가 장비 대비 ~5-10% 오차 → 임상 스크리닝에 충분'),
    ('정상범위: ', '현재 성인 일반 기준 → 연령/성별별 세분화 DB 구축 예정'),
    ('프레임률: ', '30fps 이상 권장 → 대부분 스마트폰이 지원하므로 실사용 문제 없음'),
]
for bold, text in limits:
    add_bullet(text, bold_prefix=bold)

# ═══════════════════════════════════════════════════
# 8. 활용 시나리오
# ═══════════════════════════════════════════════════
doc.add_heading('8. 활용 시나리오', level=1)

doc.add_heading('시나리오 1: 뇌졸중 재활', level=2)
p = doc.add_paragraph('62세 남성, 좌측 편마비. 매주 보행분석 실시.')
p.runs[0].italic = True
add_bullet('1주차: 속도 0.45 m/s (Limited Community), L/R SI 28%')
add_bullet('4주차: 속도 0.72 m/s (Community Amb.), L/R SI 12%')
add_bullet('8주차: 속도 0.95 m/s (Full Community), L/R SI 7%')
p = doc.add_paragraph()
run = p.add_run('→ 추이 그래프로 치료 효과 객관적 입증 → 보험 심사 자료로 활용')
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_heading('시나리오 2: 요양시설 낙상 예방', level=2)
p = doc.add_paragraph('80세 여성, 최근 보행 불안정 호소.')
p.runs[0].italic = True
add_bullet('분석 결과: 속도 0.35 m/s (Household), 안정성 보상 패턴 감지')
add_bullet('자동 알림: "낙상 위험 1.5배 증가 (Verghese 2009)"')
p = doc.add_paragraph()
run = p.add_run('→ 조기 중재 및 낙상 예방 프로그램 즉시 적용')
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_heading('시나리오 3: 무릎 치환술 후 복귀 평가', level=2)
p = doc.add_paragraph('68세 여성, 우측 TKR 수술 후 6주.')
p.runs[0].italic = True
add_bullet('분석 결과: 우측 보폭 42cm (정상 대비 ↓40%), 좌우 SI 22%')
add_bullet('패턴: 단보 패턴 + 우측 유각기 감소')
p = doc.add_paragraph()
run = p.add_run('→ 우측 하지 추진력 강화 운동 처방의 근거 자료')
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ═══════════════════════════════════════════════════
# 9. 로드맵
# ═══════════════════════════════════════════════════
doc.add_heading('9. 로드맵', level=1)
add_table(
    ['단계', '내용', '상태'],
    [
        ['v1.0', '10MWT AI 분석 + 14가지 지표 + 판정 시스템', '완료'],
        ['v1.1', '환자 관리 + 이력 추적 + 검사 비교', '완료'],
        ['v1.2', '영상 근거 클립 + 분석 플레이어', '완료'],
        ['v2.0', '연령/성별별 정상범위 DB', '개발 예정'],
        ['v2.1', 'PDF 리포트 자동 생성 + 출력', '개발 예정'],
        ['v2.2', '다중 카메라 앵글 지원', '연구 단계'],
        ['v3.0', '클라우드 SaaS + 모바일 앱', '기획 단계'],
    ],
    [0.8, 3.5, 1.2]
)

# ═══════════════════════════════════════════════════
# 10. 참고문헌
# ═══════════════════════════════════════════════════
doc.add_heading('10. 참고문헌', level=1)
refs = [
    'Perry J, et al. (1995). Classification of walking handicap in the stroke population. Stroke, 26(6), 982-989.',
    'Patterson KK, et al. (2010). Gait asymmetry in community-ambulating stroke survivors. Archives of Physical Medicine and Rehabilitation, 91(6), 884-891.',
    'Herzog W, et al. (1989). Asymmetries in ground reaction force patterns in normal human gait. Medicine and Science in Sports and Exercise, 21(1), 110-114.',
    'Verghese J, et al. (2009). Quantitative gait dysfunction and risk of cognitive decline and dementia. Journal of Neurology, Neurosurgery & Psychiatry, 80(5), 580-585.',
    'Cesari M, et al. (2005). Prognostic value of usual gait speed in well-functioning older people. Journal of the American Geriatrics Society, 53(10), 1675-1680.',
    'AAFP (2010). Gait Disorders in Older Adults. American Family Physician, 82(1), 61-68.',
    'Kim YH, et al. (2024). Gait characteristics and fall risk in Korean elderly population. Journal of Korean Physical Therapy.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'{i}. {ref}')
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(9)

# ── 저장 ──
output_path = os.path.join(os.path.dirname(__file__), 'AI_보행분석_시스템_PRD.docx')
doc.save(output_path)
print(f'Word 파일 생성 완료: {output_path}')
